from docx import Document
from docx.oxml import OxmlElement
from copy import deepcopy
from typing import List, Dict, Tuple
from src.detectors import PIIMatch

class DocxHandler:
    """Handle reading/writing docx files with PII redaction."""
    
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.doc = Document(filepath)
    
    def extract_text_with_positions(self) -> str:
        """Extract all text from document, preserving structure info."""
        text_parts = []
        
        # Extract from paragraphs
        for para in self.doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    
    def apply_redactions(self, replacements: Dict[str, str]) -> None:
        """Apply PII redactions to the document."""
        # Replace in paragraphs
        for para in self.doc.paragraphs:
            self._replace_in_paragraph(para, replacements)
        
        # Replace in tables
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)
    
    def _replace_in_paragraph(self, paragraph, replacements: Dict[str, str]) -> None:
        """Replace text in a paragraph while preserving formatting."""
        # Get full paragraph text
        full_text = paragraph.text
        
        # Check if any replacements apply
        applicable = {orig: repl for orig, repl in replacements.items() if orig in full_text}
        
        if not applicable:
            return
        
        # Sort by length (longest first) to avoid partial replacements
        sorted_replacements = sorted(applicable.items(), key=lambda x: len(x[0]), reverse=True)
        
        # Clear paragraph and rebuild with replacements
        # We need to be careful to preserve formatting
        for orig, replacement in sorted_replacements:
            if orig in full_text:
                full_text = full_text.replace(orig, replacement)
        
        # Clear existing runs and create new ones
        for run in paragraph.runs:
            r = run._element
            r.getparent().remove(r)
        
        # Add new run with replaced text
        paragraph.add_run(full_text)
    
    def save(self, output_path: str) -> None:
        """Save the document to a file."""
        self.doc.save(output_path)
    
    def save_copy(self, output_path: str) -> None:
        """Save a copy of the original document."""
        original_doc = Document(self.filepath)
        original_doc.save(output_path)

class DocumentAnalyzer:
    """Analyze document structure and content."""
    
    def __init__(self, doc: Document):
        self.doc = doc
    
    def get_paragraphs_count(self) -> int:
        return len(self.doc.paragraphs)
    
    def get_tables_count(self) -> int:
        return len(self.doc.tables)
    
    def get_total_text_length(self) -> int:
        """Get total character count of document."""
        total = 0
        for para in self.doc.paragraphs:
            total += len(para.text)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    total += len(cell.text)
        return total
    
    def get_document_structure_info(self) -> Dict:
        """Get structured info about document."""
        return {
            'paragraphs': self.get_paragraphs_count(),
            'tables': self.get_tables_count(),
            'total_text_chars': self.get_total_text_length(),
        }
